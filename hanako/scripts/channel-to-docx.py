#!/usr/bin/env python3
"""
channel-to-docx.py — 将聊天频道记录导出为结构化 .docx 文档

用法：
  python channel-to-docx.py < 输入.json          # 从 stdin 读 JSON 消息
  python channel-to-docx.py -o output.docx < in.json

输入 JSON 格式（与 parseChannel 输出一致）：
  {
    "displayName": "频道名称",
    "channelType": "channel" | "dm",
    "now": "2026-07-28T12:00:00.000Z",
    "messages": [
      { "speaker": "aris", "timestamp": "2026-07-28 05:39:19", "body": "消息正文" },
      ...
    ]
  }
"""

import sys
import json
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Color palette (matches HTML export) ──────────────────────────
BG_COLOR         = "F5F0E8"   # warm paper
HEADER_BG_START  = "FAF6EF"
HEADER_BG_END    = "F0E8D8"
HEADER_BORDER    = "DDD6C8"
TEXT_COLOR       = "2C2C2C"
MUTED_COLOR      = "888888"
DATE_LINE_COLOR  = "999999"
DATE_BG          = "F5F0E8"   # same as body, for date overlay
DIVIDER_COLOR    = "DDD6C8"
CODE_BG          = "F0EBE3"

SPEAKER_COLORS = {
    "aris":   "7C5CFC",
    "hanako-2": "F59E0B",
    "hanako": "F59E0B",
    "butter": "EC4899",
    "lorry":  "10B981",
}

def get_speaker_color(name: str) -> str:
    name_lower = name.lower()
    if name_lower in SPEAKER_COLORS:
        return SPEAKER_COLORS[name_lower]
    h = 0
    for c in name:
        h = ord(c) + ((h << 5) - h)
    hue = abs(h) % 360
    # Convert HSL to hex-ish, but for simplicity return a hex from the hue range
    # Use a simple approach: just return a color based on hue
    return f"hsl({hue}, 65%, 50%)"


def parse_timestamp(ts: str) -> datetime:
    """Parse timestamp string to datetime. Supports both with and without seconds."""
    try:
        return datetime.strptime(ts, "%Y-%m-%d %H:%M:%S")
    except ValueError:
        pass
    try:
        return datetime.strptime(ts, "%Y-%m-%d %H:%M")
    except ValueError:
        pass
    # Fallback: try ISO format
    try:
        return datetime.fromisoformat(ts)
    except ValueError:
        return datetime.now()


def format_time(ts: str) -> str:
    """Format timestamp for display (Chinese locale style)."""
    try:
        d = parse_timestamp(ts)
        return d.strftime("%Y/%m/%d %H:%M")
    except:
        return ts


def format_date_group(ts: str) -> str:
    """Format the date separator label."""
    try:
        d = parse_timestamp(ts)
        weekdays = ["星期一", "星期二", "星期三", "星期四", "星期五", "星期六", "星期日"]
        wd = weekdays[d.weekday()]
        return f"{d.year}年{d.month}月{d.day}日 {wd}"
    except:
        return ts


def set_cell_shading(cell, color_hex: str):
    """Set background color of a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_shading(paragraph, color_hex: str):
    """Set background shading for a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    pPr.append(shading)


def add_date_separator(doc, date_label: str):
    """Add a date separator (like the HTML date-separator)."""
    # Thin horizontal line with centered text overlay
    # We use a table with one row and center text
    line_table = doc.add_table(rows=1, cols=1)
    line_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = line_table.rows[0].cells[0]
    cell.text = ""
    # Remove default cell margins
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Set cell width to 100%
    tbl_width = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    line_table._tbl.tblPr.append(tbl_width)
    
    # Add a horizontal rule paragraph then the date text
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date text is small, muted
    run = p.add_run(format_date_group(date_label))
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.bold = False
    
    # Add bottom border to the paragraph as the line
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="{DIVIDER_COLOR}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    
    # Add spacing around the date
    spacing = parse_xml(
        f'<w:spacing {nsdecls("w")} w:before="240" w:after="120"/>'
    )
    pPr.append(spacing)
    
    return line_table


def add_message(doc, msg: dict, prev_date: str = None):
    """Add a single message row to the document."""
    speaker = msg["speaker"]
    timestamp = msg["timestamp"]
    body = msg.get("body", "")
    
    # Get or generate speaker color
    color_hex = get_speaker_color(speaker)
    
    # Create a table row for the message (avatar | body)
    msg_table = doc.add_table(rows=1, cols=2)
    msg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set table width to page width
    tbl_width = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    msg_table._tbl.tblPr.append(tbl_width)
    
    # Remove borders from the table
    tbl_borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    msg_table._tbl.tblPr.append(tbl_borders)
    
    # Cell 0: Avatar circle (narrow column)
    avatar_cell = msg_table.rows[0].cells[0]
    # Set narrow width
    avatar_cell.width = Cm(1.2)
    # Remove margins
    tc_pr = avatar_cell._tc.get_or_add_tcPr()
    tc_mar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="0" w:type="dxa"/>'
        f'  <w:left w:w="0" w:type="dxa"/>'
        f'  <w:bottom w:w="0" w:type="dxa"/>'
        f'  <w:right w:w="0" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tc_pr.append(tc_mar)
    
    avatar_p = avatar_cell.paragraphs[0]
    avatar_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add the initial as a colored rounded-rect (simulate with colored text on colored bg)
    run = avatar_p.add_run(speaker[0].upper())
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Set cell background to speaker color
    set_cell_shading(avatar_cell, color_hex)
    
    # Cell 1: Message body
    body_cell = msg_table.rows[0].cells[1]
    body_cell.width = Cm(14.0)
    tc_pr_body = body_cell._tc.get_or_add_tcPr()
    tc_mar_body = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="30" w:type="dxa"/>'
        f'  <w:left w:w="120" w:type="dxa"/>'
        f'  <w:bottom w:w="30" w:type="dxa"/>'
        f'  <w:right w:w="120" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tc_pr_body.append(tc_mar_body)
    
    # Speaker name + timestamp line
    speaker_p = body_cell.paragraphs[0]
    speaker_p.paragraph_format.space_after = Pt(2)
    
    # Speaker name (colored, bold)
    run = speaker_p.add_run(speaker)
    run.font.size = Pt(10)
    run.font.bold = True
    # Parse color hex to RGB
    r = int(color_hex[0:2], 16) if len(color_hex) >= 6 else 0x7c
    g = int(color_hex[2:4], 16) if len(color_hex) >= 6 else 0x5c
    b = int(color_hex[4:6], 16) if len(color_hex) >= 6 else 0xfc
    run.font.color.rgb = RGBColor(r, g, b)
    
    # Timestamp (muted, smaller)
    time_str = format_time(timestamp)
    run = speaker_p.add_run(f"  {time_str}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    
    # Message body
    body_p = body_cell.add_paragraph()
    body_p.paragraph_format.space_before = Pt(0)
    body_p.paragraph_format.space_after = Pt(4)
    body_p.paragraph_format.line_spacing = Pt(18)
    
    # Split by code blocks or just render as plain text
    # For simplicity, render the full body. We'll handle inline code later.
    lines = body.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            body_p.add_run("\n").font.size = Pt(3)
        run = body_p.add_run(line)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)
        run.font.name = "Calibri"
    
    return msg_table


def create_docx(data: dict) -> Document:
    """Main function: creates the DOCX document from the data dict."""
    doc = Document()
    
    # ── Page setup ──
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    
    # ── Default style ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)
    
    # ── Set page background color (via adding a section background) ──
    # Note: python-docx doesn't natively support page background color.
    # We'll add it via XML manipulation.
    
    # ── Header section ──
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_width = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    header_table._tbl.tblPr.append(tbl_width)
    
    # Remove borders
    tbl_borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{DIVIDER_COLOR}"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    header_table._tbl.tblPr.append(tbl_borders)
    
    header_cell = header_table.rows[0].cells[0]
    set_cell_shading(header_cell, HEADER_BG_START)
    
    # Title
    title_p = header_cell.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(4)
    run = title_p.add_run(data.get("displayName", "聊天记录"))
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x3A, 0x3A, 0x3A)
    
    # Meta info
    channel_type = data.get("channelType", "channel")
    type_label = "群聊" if channel_type == "channel" else "私聊"
    msg_count = len(data.get("messages", []))
    
    now_str = data.get("now", datetime.now().isoformat())
    try:
        export_time = datetime.fromisoformat(now_str.replace("Z", "+00:00"))
        export_fmt = export_time.strftime("%Y/%m/%d %H:%M")
    except:
        export_fmt = now_str
    
    meta_p = header_cell.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_before = Pt(2)
    meta_p.paragraph_format.space_after = Pt(24)
    
    meta_text = f"{type_label}  |  {msg_count} 条消息  |  导出于 {export_fmt}"
    run = meta_p.add_run(meta_text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    
    # ── Messages ──
    messages = data.get("messages", [])
    prev_date = None
    
    for msg in messages:
        msg_date = msg["timestamp"][:10]
        
        # Add date separator if the date changes
        if msg_date != prev_date:
            if prev_date is not None:
                # Add some space before date separator
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_before = Pt(0)
                spacer.paragraph_format.space_after = Pt(0)
                run = spacer.add_run("")
                run.font.size = Pt(1)
            
            add_date_separator(doc, msg["timestamp"])
            prev_date = msg_date
        
        add_message(doc, msg, prev_date)
    
    # ── Footer ──
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(32)
    footer_p.paragraph_format.space_after = Pt(16)
    
    # Add a top border (divider line)
    pPr = footer_p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="8" w:color="{DIVIDER_COLOR}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    
    try:
        export_date = datetime.fromisoformat(now_str.replace("Z", "+00:00"))
        date_str = export_date.strftime("%Y-%m-%d")
    except:
        date_str = now_str[:10]
    
    run = footer_p.add_run(f"由 Hana 频道导出 · {date_str}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    
    return doc


def main():
    import argparse
    
    parser = argparse.ArgumentParser(description="将频道聊天记录导出为 .docx")
    parser.add_argument("-o", "--output", default="chat-export.docx",
                        help="输出 .docx 文件路径")
    parser.add_argument("input", nargs="?", default=None,
                        help="输入 JSON 文件路径（默认从 stdin 读取）")
    args = parser.parse_args()
    
    # Read JSON input
    if args.input:
        with open(args.input, "r", encoding="utf-8") as f:
            data = json.load(f)
    else:
        data = json.load(sys.stdin)
    
    doc = create_docx(data)
    doc.save(args.output)
    
    print(json.dumps({"output": args.output, "status": "ok"}))


if __name__ == "__main__":
    main()
